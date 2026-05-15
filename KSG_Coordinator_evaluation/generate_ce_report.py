import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI
from dotenv import load_dotenv

# ===== LOAD ENV =====
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ===== TABLE BORDER FUNCTION =====
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)


# ===== LLM FUNCTION =====
def analyze_qualitative(text):

    if not text.strip():
        return "No comments provided."

    # Limit excessively long prompts
    text = text[:4000]

    prompt = f"""
The following comments were provided by participants during evaluation of programme coordination and administration at the Kenya School of Government.

Participant Feedback:
{text}

Write exactly TWO concise professional paragraphs.

Paragraph 1:
Summarize the most recurring positive feedback regarding programme coordination, administration, communication, organization, participant support, and overall management of the programme.

Paragraph 2:
Summarize the most recurring suggestions for improvement. Focus only on issues mentioned by multiple participants or issues that appear significant.

Requirements:
- Use formal institutional language
- Sound human and evidence-based
- Avoid exaggerated praise
- Avoid generic AI wording
- Avoid repetition
- Do not use bullets or headings
- Keep paragraphs concise and professional
"""

    response = client.chat.completions.create(

        model="gpt-5-nano-2025-08-07",

        messages=[

            {
                "role": "system",
                "content":
                (
                    "You are an institutional monitoring and evaluation officer "
                    "writing formal Kenya School of Government evaluation reports. "
                    "Write in a professional, concise, evidence-based and human tone. "
                    "Avoid exaggerated language, repetition, and generic AI wording."
                )
            },

            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return response.choices[0].message.content.strip()


# ===== INPUT =====
file_name = input("Enter cleaned CE file: ").strip()
df = pd.read_excel(file_name)

program_title = input("Enter Program Title: ")
date_range = input("Enter Program Dates: ")

coordinator = df["Coordinator Name"].iloc[0] if "Coordinator Name" in df.columns else "N/A"

# ===== IDENTIFY RATING COLUMNS =====
rating_cols = [
    col for col in df.columns
    if df[col].dtype in ["int64", "float64", "Int64"]
]

exclude = ["Timetable No"]
rating_cols = [col for col in rating_cols if col not in exclude]

# ===== CREATE DOCUMENT =====
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ===== HEADER =====
doc.add_paragraph("KSG/17/FEF/09")
doc.add_paragraph("KENYA SCHOOL OF GOVERNMENT")
doc.add_paragraph("CAMPUS/INSTITUTE: MATUGA")
doc.add_heading("PROGRAM COORDINATOR EVALUATION REPORT", level=1)

# ===== DETAILS TABLE =====
details_table = doc.add_table(rows=3, cols=2)

details_table.cell(0, 0).text = "PROGRAM TITLE:"
details_table.cell(0, 1).text = program_title

details_table.cell(1, 0).text = "DATE:"
details_table.cell(1, 1).text = date_range

details_table.cell(2, 0).text = "COORDINATOR’S NAME:"
details_table.cell(2, 1).text = coordinator

# Bold first column
for row in details_table.rows:
    for paragraph in row.cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True

set_table_borders(details_table)

# ===== INTRO =====
doc.add_paragraph(
    "\nThis report provides information to KSG management for decision making and action. "
    "The Coordinator Evaluation forms are filled out by participants in the course of training programmes. "
    "The Head of Department - Training is expected to discuss the evaluations with individual coordinators where necessary."
)

doc.add_paragraph(
    "\n1. Ratings on the following aspects of the program coordination "
    "(Ranging from 5=Excellent, 4=Very Good, 3=Good, 2=Fair, 1=Poor)."
)

# ===== TABLE (PERCENTAGES) =====
table = doc.add_table(rows=1, cols=6)

headers = [
    "Specific Aspects",
    "Excellent % : 5",
    "Very Good % : 4",
    "Good % : 3",
    "Fair % : 2",
    "Poor % : 1"
]

for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

for col in rating_cols:

    counts = df[col].value_counts().to_dict()

    count_5 = counts.get(5, 0)
    count_4 = counts.get(4, 0)
    count_3 = counts.get(3, 0)
    count_2 = counts.get(2, 0)
    count_1 = counts.get(1, 0)

    total_valid = count_5 + count_4 + count_3 + count_2 + count_1

    if total_valid > 0:
        p5 = round((count_5 / total_valid) * 100, 1)
        p4 = round((count_4 / total_valid) * 100, 1)
        p3 = round((count_3 / total_valid) * 100, 1)
        p2 = round((count_2 / total_valid) * 100, 1)
        p1 = round((count_1 / total_valid) * 100, 1)
    else:
        p5 = p4 = p3 = p2 = p1 = 0

    row = table.add_row().cells
    row[0].text = col
    row[1].text = str(p5)
    row[2].text = str(p4)
    row[3].text = str(p3)
    row[4].text = str(p2)
    row[5].text = str(p1)

set_table_borders(table)

# ===== QUALITATIVE =====
likes = "; ".join(df["Like"].dropna().astype(str)) if "Like" in df.columns else ""
suggestions = "; ".join(df["Suggestions"].dropna().astype(str)) if "Suggestions" in df.columns else ""

combined = f"Most liked: {likes}\nSuggestions: {suggestions}"

analysis = analyze_qualitative(combined)
paras = analysis.split("\n\n")

doc.add_paragraph("\n2. MOST liked about the program coordination and overall administration")
doc.add_paragraph(paras[0] if len(paras) > 0 else "")

doc.add_paragraph("\n3. Suggestions on how the coordinator(s) can improve on these aspects.")
doc.add_paragraph(paras[1] if len(paras) > 1 else "")

# ===== HOD =====
doc.add_paragraph("\n4. Head of Department – Training’s recommendations:")
doc.add_paragraph("......................................................................................................................")

doc.add_paragraph("\nHead of Department: Name………………. Signed: ………………… Date: …………………")

# ===== SAVE =====
base_name = os.path.splitext(file_name)[0]
output_file = f"{base_name}_CE_report.docx"

doc.save(output_file)

print(f"\n✅ Final CE report generated: {output_file}")