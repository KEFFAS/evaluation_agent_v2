import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI
from dotenv import load_dotenv

# =========================================================
# LOAD ENVIRONMENT VARIABLES
# =========================================================
load_dotenv()

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

# =========================================================
# TABLE BORDER FUNCTION
# =========================================================
def set_table_borders(table):

    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in [
        'top',
        'left',
        'bottom',
        'right',
        'insideH',
        'insideV'
    ]:

        border = OxmlElement(f'w:{border_name}')

        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:color'), '000000')

        borders.append(border)

    tblPr.append(borders)

# =========================================================
# LLM FUNCTION
# =========================================================
def generate_text(prompt):

    response = client.chat.completions.create(

        model="gpt-5-nano-2025-08-07",

        messages=[

            {
                "role": "system",
                "content":
                (
                    "You are an institutional monitoring and evaluation officer "
                    "writing formal Kenya School of Government evaluation reports. "
                    "Write in a professional, concise, evidence-based and human tone. "
                    "Avoid exaggerated language, repetition, and generic AI wording."
                )
            },

            {
                "role": "user",
                "content": prompt
            }
        ]

        
    )

    return response.choices[0].message.content.strip()

# =========================================================
# INPUT FILE
# =========================================================
file_name = input("Enter cleaned EEE file: ").strip()

# =========================================================
# LOAD FILE
# =========================================================
df = pd.read_excel(file_name)

print("Loaded:", file_name)

# =========================================================
# PROGRAM DETAILS
# =========================================================
program_title = (
    df["Program Title"].iloc[0]
    if "Program Title" in df.columns
    else "N/A"
)

coordinator = (
    df["Coordinator Name"].iloc[0]
    if "Coordinator Name" in df.columns
    else "N/A"
)

program_code = (
    df["Program Code"].iloc[0]
    if "Program Code" in df.columns
    else "N/A"
)

venue = (
    df["Venue / Campus"].iloc[0]
    if "Venue / Campus" in df.columns
    else "N/A"
)

assistant = (
    df["Program Assistant Name"].iloc[0]
    if "Program Assistant Name" in df.columns
    else "N/A"
)

duration = input("Enter Program Duration: ")

# =========================================================
# DETECT NUMERIC COLUMNS
# =========================================================
rating_cols = [
    col for col in df.columns
    if df[col].dtype in ["int64", "float64", "Int64"]
]

exclude = ["Timetable No"]

rating_cols = [
    col for col in rating_cols
    if col not in exclude
]

# =========================================================
# DETECT SPECIAL COLUMNS
# =========================================================
objective_col = next(
    (
        col for col in rating_cols
        if "objective" in col.lower()
    ),
    None
)

expectation_col = next(
    (
        col for col in rating_cols
        if "expectation" in col.lower()
    ),
    None
)

comparison_col = next(
    (
        col for col in rating_cols
        if "similar institution" in col.lower()
    ),
    None
)

# =========================================================
# CREATE DOCUMENT
# =========================================================
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# =========================================================
# HEADER
# =========================================================
doc.add_paragraph("KSG/17/EOEEF/07")
doc.add_paragraph("KENYA SCHOOL OF GOVERNMENT")
doc.add_paragraph("MATUGA")
doc.add_heading("END-OF-EVENT EVALUATION REPORT", level=1)

# =========================================================
# PROGRAM DETAILS TABLE
# =========================================================
table = doc.add_table(rows=3, cols=4)

table.cell(0, 0).text = "PROGRAMME TITLE:"
table.cell(0, 1).text = str(program_title)

table.cell(0, 2).text = "DURATION:"
table.cell(0, 3).text = str(duration)

table.cell(1, 0).text = "PROGRAM CODE:"
table.cell(1, 1).text = str(program_code)

table.cell(1, 2).text = "VENUE:"
table.cell(1, 3).text = str(venue)

table.cell(2, 0).text = "COORDINATOR:"
table.cell(2, 1).text = str(coordinator)

table.cell(2, 2).text = "PROGRAM ASST:"
table.cell(2, 3).text = str(assistant)

set_table_borders(table)

# =========================================================
# INTRODUCTION
# =========================================================
doc.add_heading("A. PROGRAMME EVALUATION", level=2)

doc.add_paragraph(
    "KSG is committed to providing quality training programmes "
    "to its customers. The evaluation findings presented in this "
    "report are intended to support continuous improvement in "
    "programme design, delivery and participant experience."
)

# =========================================================
# SECTION 1
# =========================================================
doc.add_heading(
    "1. Course Objectives Achievement",
    level=2
)

table1 = doc.add_table(rows=1, cols=2)

table1.rows[0].cells[0].text = "Rating"
table1.rows[0].cells[1].text = "Percentage of Respondents"

section1_results = []

if objective_col:

    counts = df[objective_col].value_counts().to_dict()
    total = sum(counts.values())

    labels = {
        5: "Excellent",
        4: "Very Good",
        3: "Satisfactory",
        2: "Poor",
        1: "Very Poor"
    }

    for score in [5,4,3,2,1]:

        pct = (
            round((counts.get(score,0)/total)*100,1)
            if total > 0 else 0
        )

        row = table1.add_row().cells

        row[0].text = labels[score]
        row[1].text = str(pct)

        section1_results.append(
            f"{labels[score]} = {pct}%"
        )

set_table_borders(table1)

prompt = f"""
Interpret the following course objective achievement results:

{section1_results}

Write one concise institutional paragraph similar to a Kenya School of Government evaluation report.
"""

doc.add_paragraph(generate_text(prompt))

# =========================================================
# SECTION 2
# =========================================================
doc.add_heading(
    "2. Fulfilment of Personal Expectations",
    level=2
)

table2 = doc.add_table(rows=1, cols=2)

table2.rows[0].cells[0].text = "Rating"
table2.rows[0].cells[1].text = "Percentage of Respondents"

section2_results = []

if expectation_col:

    counts = df[expectation_col].value_counts().to_dict()
    total = sum(counts.values())

    labels = {
        5: "Great Extent",
        4: "Some Extent",
        3: "Satisfactory",
        2: "Not Sure",
        1: "Not at All"
    }

    for score in [5,4,3,2,1]:

        pct = (
            round((counts.get(score,0)/total)*100,1)
            if total > 0 else 0
        )

        row = table2.add_row().cells

        row[0].text = labels[score]
        row[1].text = str(pct)

        section2_results.append(
            f"{labels[score]} = {pct}%"
        )

set_table_borders(table2)

prompt = f"""
Interpret the following participant expectation fulfilment results:

{section2_results}

Write one concise institutional paragraph similar to a Kenya School of Government evaluation report.
"""

doc.add_paragraph(generate_text(prompt))

# =========================================================
# SECTION 3
# =========================================================
doc.add_heading(
    "3. Ratings on Specific Aspects of the Training Programme",
    level=2
)

table3 = doc.add_table(rows=1, cols=6)

headers = [
    "ASPECT OF THE PROGRAMME",
    "Excellent %",
    "Very Good %",
    "Satisfactory %",
    "Poor %",
    "Very Poor %"
]

for i, h in enumerate(headers):
    table3.rows[0].cells[i].text = h

section3_summary = []

for col in rating_cols:

    if col not in [
        objective_col,
        expectation_col,
        comparison_col
    ]:

        counts = df[col].value_counts().to_dict()
        total = sum(counts.values())

        excellent = round((counts.get(5,0)/total)*100,1) if total else 0
        very_good = round((counts.get(4,0)/total)*100,1) if total else 0
        satisfactory = round((counts.get(3,0)/total)*100,1) if total else 0
        poor = round((counts.get(2,0)/total)*100,1) if total else 0
        very_poor = round((counts.get(1,0)/total)*100,1) if total else 0

        row = table3.add_row().cells

        row[0].text = str(col)
        row[1].text = str(excellent)
        row[2].text = str(very_good)
        row[3].text = str(satisfactory)
        row[4].text = str(poor)
        row[5].text = str(very_poor)

        section3_summary.append(
            f"{col}: Excellent={excellent}%, Very Good={very_good}%"
        )

set_table_borders(table3)

prompt = f"""
Interpret the following programme aspect ratings:

{section3_summary}

Write one concise institutional paragraph highlighting overall participant satisfaction trends.
"""

doc.add_paragraph(generate_text(prompt))

# =========================================================
# QUALITATIVE SECTIONS
# =========================================================
qualitative_mapping = {
    "4. Suggestions on the aspects listed in (3) above.":
        "suggestions on aspects",

    "5. Areas to be added to this training programme":
        "other areas you would like added",

    "6. Interest in other KSG programmes":
        "other ksg training programs",

    "7. Interest in additional training areas not currently offered by KSG":
        "other training programs not currently offered",

    "9. General Comments":
        "other comments"
}

for section_title, keyword in qualitative_mapping.items():

    matching_cols = [
        col for col in df.columns
        if keyword in str(col).lower()
    ]

    if matching_cols:

        col = matching_cols[0]

        responses = (
            df[col]
            .dropna()
            .astype(str)
            .str.strip()
        )

        responses = responses[
            responses != ""
        ].drop_duplicates()

        joined_text = " ".join(responses)

        prompt = f"""
The following are participant responses from a Kenya School of Government training evaluation:

{joined_text}

Write one professional paragraph summarizing the most recurring themes only.
Avoid repetition and avoid listing every response individually.
"""

        doc.add_heading(section_title, level=2)

        doc.add_paragraph(
            generate_text(prompt)
        )

# =========================================================
# SECTION 8
# =========================================================
doc.add_heading(
    "8. Rating of KSG’s Training Compared to Similar Institutions",
    level=2
)

table8 = doc.add_table(rows=1, cols=2)

table8.rows[0].cells[0].text = "Rating"
table8.rows[0].cells[1].text = "Percentage of Respondents"

comparison_results = []

if comparison_col:

    counts = df[comparison_col].value_counts().to_dict()
    total = sum(counts.values())

    labels = {
        5: "Very High",
        4: "High",
        3: "Average",
        2: "Low",
        1: "Very Low"
    }

    for score in [5,4,3,2,1]:

        pct = (
            round((counts.get(score,0)/total)*100,1)
            if total > 0 else 0
        )

        row = table8.add_row().cells

        row[0].text = labels[score]
        row[1].text = str(pct)

        comparison_results.append(
            f"{labels[score]} = {pct}%"
        )

set_table_borders(table8)

prompt = f"""
Interpret the following institutional comparison ratings:

{comparison_results}

Write one concise institutional paragraph similar to a Kenya School of Government evaluation report.
"""

doc.add_paragraph(generate_text(prompt))

# =========================================================
# SECTION 10
# =========================================================
doc.add_heading(
    "10. Key Recommendations",
    level=2
)

recommendation_text = ""

for col in df.columns:

    if any(keyword in str(col).lower() for keyword in [
        "suggest",
        "comment",
        "area"
    ]):

        recommendation_text += " ".join(
            df[col]
            .dropna()
            .astype(str)
        )

prompt = f"""
Based on the following participant feedback:

{recommendation_text}

Generate concise actionable institutional recommendations for improving future programmes.
"""

doc.add_paragraph(
    generate_text(prompt)
)

# =========================================================
# SIGNATURE SECTION
# =========================================================
doc.add_paragraph(
    "\nPrepared by………..………..….………...…………"
    "Date…………….…………...Signature……………"
)

doc.add_paragraph(
    "Confirmed by………..….………...…………………"
    "Date…………………………Signature……………"
)

doc.add_paragraph(
    "Approved by………………...…..…….………..……"
    "Date…………………………Signature……………"
)

# =========================================================
# SAVE DOCUMENT
# =========================================================
base_name = os.path.splitext(file_name)[0]

output_file = f"{base_name}_EEE_Report_LLM.docx"

doc.save(output_file)

print("\n===================================")
print("✅ EEE LLM REPORT GENERATED")
print("Saved as:")
print(output_file)
print("===================================")