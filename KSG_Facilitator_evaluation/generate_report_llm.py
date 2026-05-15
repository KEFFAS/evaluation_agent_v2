import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI
from dotenv import load_dotenv

# ===== LOAD ENV =====
load_dotenv()

# ===== INIT CLIENT =====
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ===== TABLE BORDER FUNCTION =====
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)

# ===== LLM FUNCTION =====
def analyze_qualitative(text):

    if not text.strip():
        return "No comments provided."

    # Prevent excessively large prompts
    text = text[:4000]

    prompt = f"""
The following comments were provided by participants during evaluation of a facilitator at the Kenya School of Government.

Participant Feedback:
{text}

Write exactly TWO concise professional paragraphs.

Paragraph 1:
Summarize the most recurring positive feedback regarding facilitation, delivery style, subject mastery, participant engagement, communication, responsiveness, and overall teaching effectiveness.

Paragraph 2:
Summarize the most recurring suggestions for improvement. Focus only on issues mentioned repeatedly or issues that appear significant.

Requirements:
- Use formal institutional language
- Sound human and evidence-based
- Avoid exaggerated praise
- Avoid generic AI wording
- Avoid repetition
- Do not use bullet points
- Do not use headings
- Write as a flowing narrative suitable for an official evaluation report
"""

    response = client.chat.completions.create(

        model="gpt-5-nano-2025-08-07",

        messages=[

            {
                "role": "system",
                "content":
                (
                    "You are an institutional monitoring and evaluation officer "
                    "writing formal Kenya School of Government evaluation reports. "
                    "Write in a professional, concise, evidence-based and human tone. "
                    "Avoid exaggerated language, repetition, and generic AI wording."
                )
            },

            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return response.choices[0].message.content.strip()



# ===== LOAD DATA =====
file_name = input("Enter cleaned file name: ").strip()
df = pd.read_excel(file_name)

program_title = input("Enter Program Title: ")
report_date = input("Enter Report Date: ")
total_participants = int(input("Enter total participants: "))

# ===== CLEAN COLUMNS =====
df.columns = df.columns.astype(str).str.strip().str.title()

rename_map = {
    "Topic": "Topic Description",
    "Session Topic": "Topic Description",
    "Facilitator": "Lecturer Name",
    "Lecturer": "Lecturer Name"
}
df = df.rename(columns=rename_map)

# ===== SORT =====
df = df.sort_values(by=["Lecturer Name", "Topic Description"])

# ===== RATING COLUMNS =====
rating_cols = [
    "Punctuality",
    "Presentation Flow",
    "Handling Questions",
    "Active Participation Of Learners",
    "Use Of Visual Aids",
    "Relevance Of Subject To Workplace",
    "Use Of Relevant Examples",
    "Knowledge Of Subject",
    "Treats Participants With Dignity And Respect",
    "Variety And Appropriateness Of Training Methods"
]

rating_cols = [col for col in rating_cols if col in df.columns]

# ===== GROUP =====
grouped = df.groupby(["Lecturer Name", "Topic Description"])

# ===== CREATE DOCUMENT =====
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ===== LOOP =====
for (facilitator, session), group in grouped:

    # ===== HEADER =====
    doc.add_paragraph("KSG/17/FER/08")
    doc.add_paragraph("KENYA SCHOOL OF GOVERNMENT")
    doc.add_paragraph("MATUGA")
    doc.add_heading("FACILITATOR EVALUATION REPORT", level=1)

    # ===== DETAILS TABLE =====
    details_table = doc.add_table(rows=3, cols=2)

    details_table.cell(0, 0).text = "PROGRAM TITLE:"
    details_table.cell(0, 1).text = program_title

    details_table.cell(1, 0).text = "SESSION’S TOPIC:"
    details_table.cell(1, 1).text = session

    details_table.cell(2, 0).text = f"FACILITATOR: {facilitator}"
    details_table.cell(2, 1).text = f"DATE: {report_date}"

    set_table_borders(details_table)

    doc.add_paragraph("")

    # ===== DESCRIPTION =====
    doc.add_paragraph(
        "This report provides information to KSG management for decision making and action."
    )

    doc.add_paragraph("\nI. Participants’ Ratings.")

    # ===== RATINGS TABLE =====
    table = doc.add_table(rows=1, cols=10)

    headers = [
        "SPECIFIC ASPECTS", "Total no of participants", "Non response",
        "5", "4", "3", "2", "1", "Total valid responses", "% of Scores 4 &5"
    ]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for col in rating_cols:

        counts = group[col].value_counts().to_dict()

        count_5 = counts.get(5, 0)
        count_4 = counts.get(4, 0)
        count_3 = counts.get(3, 0)
        count_2 = counts.get(2, 0)
        count_1 = counts.get(1, 0)

        total_valid = count_5 + count_4 + count_3 + count_2 + count_1
        non_response = total_participants - total_valid
        percent_45 = ((count_5 + count_4) / total_valid * 100) if total_valid > 0 else 0

        row_cells = table.add_row().cells
        values = [
            col,
            total_participants,
            non_response,
            count_5,
            count_4,
            count_3,
            count_2,
            count_1,
            total_valid,
            round(percent_45, 1)
        ]

        for i, val in enumerate(values):
            row_cells[i].text = str(val)

    set_table_borders(table)

    doc.add_paragraph("")

    # ===== QUALITATIVE (LLM ENHANCED) =====
    likes_raw = "; ".join(group["Like"].dropna().astype(str))
    suggestions_raw = "; ".join(group["Suggestions"].dropna().astype(str))

    combined_text = f"""
Most Liked:
{likes_raw}

Suggestions:
{suggestions_raw}
"""

    qualitative = analyze_qualitative(combined_text)

    paragraphs = qualitative.split("\n\n")

    doc.add_paragraph("Most liked about the facilitator")
    doc.add_paragraph(paragraphs[0] if len(paragraphs) > 0 else "")

    doc.add_paragraph("\nSuggestions on areas of improvement")
    doc.add_paragraph(paragraphs[1] if len(paragraphs) > 1 else "Participants expressed minimal suggestions for improvement.")

    # ===== HOD =====
    doc.add_paragraph("\nII. Head of Department – Training’s comments:")
    doc.add_paragraph("................................................")

    doc.add_paragraph("\nIII. Head of Department – Training’s proposals or recommendations:")
    doc.add_paragraph("................................................")

    doc.add_page_break()

# ===== SAVE =====
base_name = os.path.splitext(file_name)[0]
output_file = f"{base_name}_KSG_report_LLM.docx"

doc.save(output_file)

print(f"\n✅ Final LLM-enhanced report generated: {output_file}")